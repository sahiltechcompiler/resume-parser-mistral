import os
import pdfplumber
import warnings
warnings.filterwarnings("ignore", message="CropBox missing from /Page")
from docx import Document
from PIL import Image
import pytesseract
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
from peft import PeftModel
import re
import json

# === Load tokenizer and LoRA-adapted model ===
model_path = "mistral_resume_parser_model"
base_model_id = "mistralai/Mistral-7B-Instruct-v0.1"

tokenizer = AutoTokenizer.from_pretrained(model_path)
base_model = AutoModelForCausalLM.from_pretrained(
    base_model_id,
    device_map="auto",
    torch_dtype=torch.float16,
    trust_remote_code=True
)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model = PeftModel.from_pretrained(base_model, model_path)
model.to(device)
model.eval()


def extract_emails(text):
    return re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', text)

def extract_phones(text):
    rough = re.findall(r'\b[+\d][\d\-\s()]{8,}\d\b', text)
    phones = []
    for m in rough:
        digits = re.sub(r'\D', '', m)
        if len(digits) >= 10:
            phone = ('+' if m.strip().startswith('+') else '') + digits
            phones.append(phone)
    return list(dict.fromkeys(phones))

def parse_resume(file_path: str) -> dict:
    text = clean_resume_text(extract_text(file_path))
    if not text:
        return {"error": "Could not extract text from resume"}

    prompt = f"<|startoftext|>\nInput: {text}\nOutput:"
    inputs = tokenizer(prompt, return_tensors="pt", padding=True, truncation=True, max_length=4096).to(device)

    input_token_count = inputs["input_ids"].shape[1]
    print(f"?? Input Tokens: {input_token_count}")

    with torch.inference_mode():
        outputs = model.generate(
            **inputs,
            max_new_tokens=2048,
            do_sample=False,
            pad_token_id=tokenizer.eos_token_id,
            eos_token_id=tokenizer.eos_token_id
        )

    output_token_count = outputs.shape[1] - input_token_count
    print(f"?? Output Tokens: {output_token_count}")
    print(f"?? Total Tokens: {outputs.shape[1]}")

    result = tokenizer.decode(outputs[0], skip_special_tokens=True)
    if "Output:" in result:
        json_part = result.split("Output:", 1)[1].split("<|endoftext|>")[0].strip()
        if json_part.startswith('"') and json_part.endswith('"'):
            json_part = json_part.encode().decode('unicode_escape').strip('"')
        try:
            parsed = json.loads(json_part)
        except json.JSONDecodeError:
            return {"error": "Generated output is not valid JSON", "raw": json_part}

        parsed["emails"] = extract_emails(text)
        parsed["phone_numbers"] = extract_phones(text)
        return parsed

    return {"error": "Model did not return output"}

def extract_text(file_path):
    ext = file_path.lower()
    if ext.endswith(".pdf"):
        full_text = ""
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text(x_tolerance=1, y_tolerance=1)
                tables = page.extract_tables()
                tables_text = ""
                for table in tables:
                    if not table or len(table) > 30:
                        continue
                    for row in table:
                        if not row:
                            continue
                        row_cleaned = [cell.strip() for cell in row if cell and cell.strip()]
                        if len(row_cleaned) < 2 or len(row_cleaned) > 8:
                            continue
                        if all(re.match(r"(?i)column \d+", cell) for cell in row_cleaned):
                            continue
                        tables_text += " | ".join(row_cleaned) + "\n"

                words = page.extract_words()
                left_column = []
                right_column = []
                if words:
                    midpoint = page.width / 2
                    for word in words:
                        if word["x0"] < midpoint:
                            left_column.append(word["text"])
                        else:
                            right_column.append(word["text"])
                if left_column and right_column:
                    merged = left_column + right_column
                    full_text += f"--- PAGE {page_num+1} ---\n" + " ".join(merged) + "\n\n"
                elif page_text:
                    full_text += f"--- PAGE {page_num+1} ---\n{page_text}\n"
                if tables_text.strip():
                    full_text += f"--- TABLES PAGE {page_num+1} ---\n{tables_text}\n"
        return full_text.strip()

    elif ext.endswith(".docx"):
        doc = Document(file_path)
        paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables_text.append(" | ".join(cells))
        full_text = "\n".join(paragraphs_text)
        if tables_text:
            full_text += "\n--- TABLES ---\n" + "\n".join(tables_text)
        return full_text.strip()

    elif ext.endswith(".txt"):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"?? Could not read .txt file: {e}")
        return ""

    elif ext.endswith((".png", ".jpg", ".jpeg")):
        try:
            image = Image.open(file_path)
            return pytesseract.image_to_string(image).strip()
        except Exception as e:
            print(f"?? Error extracting text from image {file_path}: {e}")
        return ""

    return ""

def clean_resume_text(text):
    text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text)
    text = re.sub(r"(\|[^\n]*){10,}", "", text)
    text = re.sub(r"(\|\s*){3,}", "|", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s{4,}", " ", text)
    return text.strip()
